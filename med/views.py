from urllib import request
from django.shortcuts import render
from django.shortcuts import render
from django.http import JsonResponse,HttpResponse
from .models import Patient
from django.views.decorators.csrf import csrf_exempt
import json
from datetime import datetime
import base64
from docx import Document
from docx.shared import Inches
from io import BytesIO
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Create your views here.


def login(request):
    return render(request, 'login.html')

def profile(request):
    return render(request, 'profile.html')

def signup(request):
    return render(request,'signup.html')

def popupform(request):
    return render(request,'popupform.html')
# def report(request, id):
#     return render(request, 'report.html', {'id': id})

from django.shortcuts import render, get_object_or_404
from .models import UserAccount

from django.shortcuts import render, get_object_or_404, redirect
from .models import UserAccount

def user_detail(request, id):
    user = get_object_or_404(UserAccount, id=id)

    if request.method == "POST":
        # Login Details
        user.is_active = request.POST.get('status') == 'ACTIVE'
        # user.modality = request.POST.get('modality', user.modality)

        # Personal details
        user.name = request.POST.get('personName', user.name)
        # user.contact = request.POST.get('mobileNo', user.contact)
        # user.email = request.POST.get('email', user.email)

        # Password update
        new_password = request.POST.get('password')
        if new_password:
            user.password = new_password  # Optional: hash it in production

        user.save()
        return redirect('signup')  # Refresh page after saving

    return render(request, 'user_detail.html', {'user': user})


# views.py
from django.shortcuts import redirect

def update_user(request, id):
    user = get_object_or_404(UserAccount, id=id)
    if request.method == 'POST':
        user.name = request.POST.get('accountName')
        # user.contact = request.POST.get('mobileNo')
        user.usertype = request.POST.get('userType')
        user.is_active = request.POST.get('status') == 'ACTIVE'
        # add other fields
        user.save()
        return redirect('user_detail', id=user.id)
    return render(request, 'user_detail.html', {'user': user})


# def imagingA(request):
#     patients = Patient.objects.all().order_by('-entry_time')
#     return render(request, 'imagingA.html', {'patients': patients})

# def RADS(request):
#     patients = Patient.objects.all().order_by('-entry_time')
#     return render(request, 'RADS.html', {'patients': patients})

def invoice(request):
    return render(request, 'invoice.html')
def payment(request):
    return render(request, 'payment.html')


# @csrf_exempt
# def add_patient(request):
#     if request.method == 'POST':
#         try:
#             data = json.loads(request.body)
#             print("Received data:", data)  # For debugging
            
#             # Clean up base64 image data if needed
#             scan_image = data.get('scan', '')  # Use get() with default value
#             if ',' in scan_image:
#                 scan_image = scan_image.split(',')[1]
#             center_name = data.get('center') or request.session.get('user_name')  # ✅ from frontend or session

            
#             patient = Patient.objects.create(
#                 name=data.get('name'),
#                 age=data.get('age'),
#                 gender=data.get('gender'),
#                 history=data.get('history'),
#                 scan_type=data.get('scanType'),  # Changed from scan_type
#                 body_part=data.get('bodyPart'),  # Changed from body_part
#                 ref_by=data.get('refBy'),        # Changed from ref_by
#                 scan_image=scan_image,
#                 center=center_name  # ✅ store the center name
#             )
#             return JsonResponse({
#                 'status': 'success',
#                 'patient_id': patient.patient_id
#             })
#         except KeyError as e:
#             print("KeyError:", e)  # Add debug print
#             return JsonResponse({
#                 'status': 'error',
#                 'message': f'Missing field: {str(e)}'
#             })
#         except Exception as e:
#             print("Exception:", e)  # Add debug print
#             return JsonResponse({
#                 'status': 'error',
#                 'message': str(e)
#             })
@csrf_exempt
def add_patient(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            scan_image = data.get('scan', '')
            if ',' in scan_image:
                scan_image = scan_image.split(',')[1]

            user_id = request.session.get('user_id')  # ✅ session से current user ID
            center_name = data.get('center') or request.session.get('user_name')

            patient = Patient.objects.create(
                name=data.get('name'),
                age=data.get('age'),
                gender=data.get('gender'),
                history=data.get('history'),
                scan_type=data.get('scanType'),
                body_part=data.get('bodyPart'),
                ref_by=data.get('refBy'),
                scan_image=scan_image,
                center=center_name,
                created_by_id=user_id  # ✅ लिंक हो गया user से
            )

            return JsonResponse({'status': 'success', 'patient_id': patient.patient_id})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)})

from django.shortcuts import render
from .models import Patient,UserAccount

def index(request):
    centers = UserAccount.objects.filter(usertype='IMAGING', is_active=True)
    rads = UserAccount.objects.filter(usertype='RADS', is_active=True)

    patients = Patient.objects.all().order_by('-entry_time')

    # ✅ sab unique center names nikal lo (jo blank na ho)
    # centers_qs = Patient.objects.values_list('center', flat=True).distinct().order_by('center')
    # centers = [c for c in centers_qs if c and c.strip()]

    # ✅ dono bhej do template ko
    return render(request, 'index.html', {
        'patients': patients,
        'centers': centers,
        'rads': rads
    })

@csrf_exempt
def update_report(request, id):
    if request.method == 'POST':
        data = json.loads(request.body)
        try:
            patient = Patient.objects.get(id=id)
            patient.report = data['report']
            patient.status = data['status']
            patient.save()
            return JsonResponse({'status': 'success'})
        except Patient.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Patient not found'})

def report(request, id):
    try:
        patient = Patient.objects.get(id=id)
        return render(request, 'report.html', {'patient': patient})
    except Patient.DoesNotExist:
        return render(request, 'report.html', {'error': 'Patient not found'})
    

    
@csrf_exempt
def get_patient(request, id):
    try:
        patient = Patient.objects.get(id=id)
        data = {
            'id': patient.id,
            'patient_id': patient.patient_id,
            'name': patient.name,
            'age': patient.age,
            'gender': patient.gender,
            'history': patient.history,
            'scan_type': patient.scan_type,
            'body_part': patient.body_part,
            'ref_by': patient.ref_by,
            'scan_image': patient.scan_image,
            'status': patient.status,
            'report': patient.report
        }
        return JsonResponse(data)
    except Patient.DoesNotExist:
        return JsonResponse({'error': 'Patient not found'}, status=404)





from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from .models import Patient  # apne model ka naam use karo

# @csrf_exempt
# def delete_patient(request, id):
#     if request.method == 'DELETE':
#         try:
#             patient = Patient.objects.get(id=id)
#             patient.delete()
#             return JsonResponse({'success': True})
#         except Patient.DoesNotExist:
#             return JsonResponse({'error': 'Patient not found'}, status=404)
#     return JsonResponse({'error': 'Invalid request'}, status=400)


        
        
@csrf_exempt
def download_report(request, id, format):
    try:
        patient = Patient.objects.get(id=id)
        if patient.status != 'FINAL':
            return JsonResponse({'error': 'Report not finalized'}, status=400)

        if format == 'pdf':
            # Create PDF
            buffer = BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)
            
            # Add content
            p.setFont("Helvetica", 12)
            p.drawString(100, 750, f"Patient ID: {patient.patient_id}")
            p.drawString(100, 730, f"Name: {patient.name}")
            p.drawString(100, 710, f"Age/Gender: {patient.age}/{patient.gender}")
            p.drawString(100, 690, f"Scan Type: {patient.scan_type}")
            p.drawString(100, 670, f"Body Part: {patient.body_part}")
            p.drawString(100, 650, f"Referred By: {patient.ref_by}")
            p.drawString(100, 630, "Report:")
            p.drawString(120, 610, patient.report or "No report available")
            
            p.save()
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="patient_report_{patient.patient_id}.pdf"'
            return response

        elif format == 'docx':
            # Create Word document
            doc = Document()
            doc.add_heading('Patient Report', 0)
            
            # Add content
            doc.add_paragraph(f'Patient ID: {patient.patient_id}')
            doc.add_paragraph(f'Name: {patient.name}')
            doc.add_paragraph(f'Age/Gender: {patient.age}/{patient.gender}')
            doc.add_paragraph(f'Scan Type: {patient.scan_type}')
            doc.add_paragraph(f'Body Part: {patient.body_part}')
            doc.add_paragraph(f'Referred By: {patient.ref_by}')
            doc.add_heading('Report:', level=1)
            doc.add_paragraph(patient.report or "No report available")
            
            # Save document
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="patient_report_{patient.patient_id}.docx"'
            return response

        return JsonResponse({'error': 'Invalid format'}, status=400)
        
    except Patient.DoesNotExist:
        return JsonResponse({'error': 'Patient not found'}, status=404)
    
    
from django.shortcuts import render, redirect
from .models import UserAccount
from .forms import SignupForm, LoginForm

def signup_view(request):
    form = SignupForm(request.POST or None)
    if form.is_valid():
        form.save()
        return redirect('signup')
    users_list = UserAccount.objects.all()
    return render(request, 'signup.html', {'form': form, 'users': users_list})

def login_view(request):
    form = LoginForm(request.POST or None)
    error = None
    if form.is_valid():
        userid = form.cleaned_data['userid']
        password = form.cleaned_data['password']

         # 🔹 Step 1: Default admin login check
        if userid == "Admin" and password == "12345":
            return redirect('index')  # default page after admin login
        if userid == "Superadmin" and password == "Aman@8865":
            return redirect('super_admin')  # default page after admin login
        try:
            user = UserAccount.objects.get(userid=userid, password=password)
            # ✅ Save user info to session
            request.session['user_id'] = user.id
            request.session['user_name'] = user.name
            request.session['user_type'] = user.usertype

            if user.usertype == 'SUPERADMIN':
                return redirect('super_admin')
            elif user.usertype == 'ADMIN':
                return redirect('index')
            elif user.usertype == 'IMAGING':
                return redirect('imagingA')
            elif user.usertype == 'RADS':
                return redirect('RADS')
            else:
                return redirect('login')

        except UserAccount.DoesNotExist:
            error = "Invalid userid or password"

    return render(request, 'login.html', {'form': form, 'error': error})

def imagingA(request):
    if 'user_id' not in request.session:
        return redirect('login')

    user_id = request.session['user_id']
    user_name = request.session['user_name']

    # ✅ Admin ko sab dikhayenge
    if user_name.lower() == "admin":
        patients = Patient.objects.all().order_by('-entry_time')
    else:
        patients = Patient.objects.filter(created_by_id=user_id).order_by('-entry_time')

    return render(request, 'imagingA.html', {'patients': patients})

def RADS(request):
    if 'user_id' not in request.session:
        return redirect('login')

    user_id = request.session['user_id']
    user_name = request.session['user_name']

    # Agar admin ho toh sab dikhana
    if user_name.lower() == "admin":
        patients = Patient.objects.all().order_by('-entry_time')
    else:
        # Agar current logged-in user RAD type hai to assigned patients dikhao
        user = UserAccount.objects.get(id=user_id)
        if user.usertype == 'RADS':
            patients = Patient.objects.filter(assigned_to_id=user_id).order_by('-entry_time')
        else:
            # agar imaging center user hai to apne created patients dikhao (jaise pehle)
            patients = Patient.objects.filter(created_by_id=user_id).order_by('-entry_time')

    return render(request, 'RADS.html', {'patients': patients})


def logout_view(request):
    request.session.flush()  # session clear
    return redirect('login')

# def rads_page(request):
#     return render(request, 'RADS.html')

# def imaging_page(request):
#     return render(request, 'imagingA.html')




    # views.py (kisi jagah imports ke upar)
from django.http import JsonResponse, HttpResponseForbidden, HttpResponseBadRequest
from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt  # agar aap csrf_token use nahi kar rahe toh rakh sakte ho

# Add this function somewhere in views.py
@csrf_exempt   # agar aap template me CSRF token bhejna nahi chahte toh isko rakh lo; production me hata ke proper CSRF use karo
@require_POST
def assign_patient(request, patient_id):
    print("Session check ->", request.session.get('user_name'), request.session.get('user_type'))

    """
    Frontend se AJAX POST karega: body = {"rad_id": <user_id>}
    Sirf admin ya imaging user assign kar sake — aap yahan permission logic change kar sakte ho.
    """
    # basic permission: agar session me user_name nahi hai -> not allowed
    if 'user_id' not in request.session:
        return HttpResponseForbidden("Login required")

    # Optional: sirf admin ya IMAGING type wale assign kar sake:
    user_name = request.session.get('user_name', '').lower()
    user_type = request.session.get('user_type', '')

    # allow if admin or IMAGING user
    if not (user_name == 'admin' or user_type in ['IMAGING', 'RADS']):
        return HttpResponseForbidden("Not allowed to assign")

    try:
        data = json.loads(request.body.decode('utf-8'))
    except Exception:
        return HttpResponseBadRequest("Invalid JSON")

    rad_id = data.get('rad_id')

    try:
        patient = Patient.objects.get(id=patient_id)
    except Patient.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'Patient not found'}, status=404)

    # unassign
    if not rad_id:
        patient.assigned_to = None
        patient.save()
        return JsonResponse({'success': True, 'assigned_to': None})

    try:
        rad_user = UserAccount.objects.get(id=rad_id, is_active=True, usertype='RADS')
    except UserAccount.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'RAD user not found'}, status=404)

    patient.assigned_to = rad_user
    patient.save()

    # Optional: yahan notification/email dal sakte ho

    return JsonResponse({
        'success': True,
        'assigned_to': {'id': rad_user.id, 'name': rad_user.name}
    })


from django.contrib import messages
from django.shortcuts import redirect
from django.contrib.auth.hashers import make_password  # optional, recommended
from .models import UserAccount
def super_admin(request):
    # if request.session.get('user_name', '').lower() != "admin":
    #     return HttpResponse("Access Denied")

    users = UserAccount.objects.all().order_by('-id')
    patients = Patient.objects.all().order_by('-entry_time')
    centers = UserAccount.objects.filter(usertype='IMAGING', is_active=True)
    admins = UserAccount.objects.filter(usertype='ADMIN', is_active=True)

    return render(request, 'super_admin.html', {
        'users': users,
        'patients': patients,
        'rads': UserAccount.objects.filter(usertype='RADS', is_active=True),
        'centers': centers,
        'admins': admins,
    })

    

def toggle_user_status(request, id):
    user = UserAccount.objects.get(id=id)
    user.is_active = not user.is_active
    user.save()
    return redirect('super_admin')


def change_user_role(request, id, role):
    user = UserAccount.objects.get(id=id)
    if role in ['RADS', 'IMAGING']:
        user.usertype = role
        user.save()
    return redirect('super_admin')

from django.contrib import messages
from django.shortcuts import redirect
from .models import Patient

def assign_patient_superadmin(request):
    if request.method == "POST":
        patient_id = request.POST.get("patient_id")
        rads = request.POST.get("rads")

        try:
            patient = Patient.objects.get(id=patient_id)

            # ✅ Sirf RADS assign karo
            if rads:
                patient.assigned_to_id = rads
                # patient.status = "ASSIGNED"
                patient.save()

            messages.success(request, "Patient successfully assigned to RADS.")
        except Patient.DoesNotExist:
            messages.error(request, "Patient not found.")
        except Exception as e:
            messages.error(request, f"Assignment failed: {str(e)}")

    return redirect('super_admin')



from .models import Patient, UserAccount  # example model names

def patients(request):
    patients = Patient.objects.all().order_by('-id')
    rads = UserAccount.objects.filter(usertype='rads')  # or as per your model field
    return render(request, 'patients.html', {'patients': patients, 'rads': rads})

from django.views.decorators.http import require_POST
from django.views.decorators.csrf import csrf_exempt
from django.contrib import messages
from django.http import JsonResponse
from .models import UserAccount

@csrf_exempt
@require_POST
def add_user(request):
    name = request.POST.get('name')
    userid = request.POST.get('userid')
    password = request.POST.get('password')
    usertype = request.POST.get('usertype')
    parent_admin_id = request.POST.get('parent_admin_id')

    if not (name and userid and password and usertype):
        messages.error(request, "सभी फ़ील्ड भरें!")
        return redirect('super_admin')

    if UserAccount.objects.filter(userid=userid).exists():
        messages.error(request, "यह UserID पहले से मौजूद है!")
        return redirect('super_admin')

    user = UserAccount(name=name, userid=userid, password=password, usertype=usertype, is_active=True)

    if usertype in ['IMAGING', 'RADS'] and parent_admin_id:
        try:
            parent_admin = UserAccount.objects.get(id=parent_admin_id, usertype='ADMIN')
            user.parent_admin = parent_admin
        except UserAccount.DoesNotExist:
            messages.warning(request, "Admin नहीं मिला, user बिना parent admin के बना दिया गया।")

    user.save()
    messages.success(request, f"User '{name}' सफलतापूर्वक बना दिया गया ✅")
    return redirect('super_admin')


from django.shortcuts import render, get_object_or_404
from .models import UserAccount

def admin_details(request, admin_id):
    admin_user = get_object_or_404(UserAccount, id=admin_id, usertype='ADMIN')

    imaging_users = UserAccount.objects.filter(usertype='IMAGING', parent_admin_id=admin_id)
    rads_users = UserAccount.objects.filter(usertype='RADS', parent_admin_id=admin_id)

    return render(request, 'admin_details.html', {
        'admin': admin_user,
        'imaging_users': imaging_users,
        'rads_users': rads_users,
    })


def user_list(request):
    # 🔹 सिर्फ ADMIN वाले users लाएँ
    users = UserAccount.objects.filter(usertype='ADMIN').order_by('-id')
    return render(request, 'user_list.html', {'users': users})


from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import UserAccount

def user_details_api(request, user_id):
    user = get_object_or_404(UserAccount, id=user_id)
    return JsonResponse({
        'name': user.name,
        'userid': user.userid,
        # 'email': user.email,
        # 'contact': user.contact,
        # 'modality': user.modality,
        'status': 'Active' if user.is_active else 'Inactive',
        'password': user.password,  # ⚠️ सिर्फ development के लिए
    })
